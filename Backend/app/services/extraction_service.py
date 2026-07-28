"""
Extracts raw text from uploaded contracts, page by page.

Kept separate from OCR: this service only pulls text that's already
digitally embedded in the file. Pages with little/no extractable text are
flagged so ProcessingService knows to route them through OCR instead.
"""

from dataclasses import dataclass
from pathlib import Path

import fitz  # PyMuPDF
from docx import Document as DocxDocument

# A page with fewer than this many extracted characters is treated as
# "scanned" (image-only) rather than genuinely blank - genuinely blank
# pages are rare in real contracts, so this threshold errs toward OCR.
MIN_CHARS_FOR_DIGITAL_TEXT = 20


@dataclass
class ExtractedPage:
    page_number: int  # 1-indexed
    text: str
    needs_ocr: bool


class ExtractionService:
    def extract(self, file_path: Path, mime_type: str) -> list[ExtractedPage]:
        if mime_type == "application/pdf":
            return self._extract_pdf(file_path)
        elif mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            return self._extract_docx(file_path)
        else:
            raise ValueError(f"Unsupported mime type for extraction: {mime_type}")

    def _extract_pdf(self, file_path: Path) -> list[ExtractedPage]:
        pages: list[ExtractedPage] = []
        with fitz.open(file_path) as pdf:
            for index, page in enumerate(pdf, start=1):
                text = page.get_text().strip()
                pages.append(
                    ExtractedPage(
                        page_number=index,
                        text=text,
                        needs_ocr=len(text) < MIN_CHARS_FOR_DIGITAL_TEXT,
                    )
                )
        return pages

    def _extract_docx(self, file_path: Path) -> list[ExtractedPage]:
        # DOCX has no native "page" concept (pagination is a rendering
        # detail, not stored in the file) - the whole document is treated
        # as a single logical page for chunking purposes.
        doc = DocxDocument(str(file_path))
        parts = [p.text for p in doc.paragraphs if p.text.strip()]

        # python-docx's `.paragraphs` doesn't include table content, so we
        # walk tables separately - contracts frequently put payment
        # schedules, party details, etc. in tables.
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        parts.append(cell.text.strip())

        text = "\n".join(parts).strip()
        return [ExtractedPage(page_number=1, text=text, needs_ocr=False)]